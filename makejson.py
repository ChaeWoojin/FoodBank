# Parse all tables into a dictionary keyed by heading
from collections import defaultdict
from docx import Document
import re
import json

# Load the document
doc_path = "./기술문서_코드정보별첨.docx"
doc = Document(doc_path)


parsed_data = defaultdict(list)
current_key = None

for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith("□"):
        current_key = text.replace("□", "").strip()
    elif current_key and text:
        # Try to parse lines with a format like "01 중앙" or "10 긴급지원대상자"
        match = re.match(r"^(\S+)\s+(.+)$", text)
        if match:
            code, name = match.groups()
            parsed_data[current_key].append({"code": code, "name": name})

# Also include all tables from the document if any exist
for table in doc.tables:
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    for row in table.rows[1:]:
        row_data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells) if i < len(headers)}
        if current_key:
            parsed_data[current_key].append(row_data)

# Save the entire parsed data into a JSON file
full_json_output_path = "./기술문서_전체코드정보.json"
with open(full_json_output_path, "w", encoding="utf-8") as f:
    json.dump(parsed_data, f, ensure_ascii=False, indent=2)

full_json_output_path
